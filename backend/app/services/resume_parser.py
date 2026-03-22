"""
Resume Parser Service - Extract and structure resume text into JSON Resume schema
Handles PDF and DOCX extraction with intelligent section detection
"""

import re
from typing import Optional, List, Tuple
from datetime import datetime

from app.services.resume_schema import (
    ResumeSchema,
    ResumeBasics,
    WorkExperience,
    Education,
    Skill,
    Certificate,
    Project,
    Location,
    Profile,
)


class ResumeParser:
    """
    Parse raw resume text into structured JSON Resume format.
    Uses pattern matching and heuristics for section detection.
    """

    SECTION_PATTERNS = {
        "contact": [
            r"(?:email|e-mail)[:\s]*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})",
            r"phone[:\s]*([+]?[\d\s\-\(\)]{10,})",
            r"location[:\s]*([^\n]+)",
            r"linkedin[:\s]*(https?://[^\s]+)",
            r"github[:\s]*(https?://[^\s]+)",
        ],
        "summary": [
            r"(?:summary|profile|objective|about)[:\s]*\n(.*?)(?=\n[A-Z]|\n\n|\Z)",
        ],
        "experience": [
            r"(?:experience|work|employment|professional experience)[:\s]*\n(.*?)(?=\n[A-Z]|\n\n|\Z)",
        ],
        "education": [
            r"(?:education|academic|qualification)[:\s]*\n(.*?)(?=\n[A-Z]|\n\n|\Z)",
        ],
        "skills": [
            r"(?:skills|technologies|technical skills|competencies)[:\s]*\n(.*?)(?=\n[A-Z]|\n\n|\Z)",
        ],
        "projects": [
            r"(?:projects|portfolio)[:\s]*\n(.*?)(?=\n[A-Z]|\n\n|\Z)",
        ],
        "certifications": [
            r"(?:certifications|certificates|credentials)[:\s]*\n(.*?)(?=\n[A-Z]|\n\n|\Z)",
        ],
    }

    def __init__(self, raw_text: str):
        self.raw_text = raw_text
        self.cleaned_text = self._clean_text(raw_text)

    def _clean_text(self, text: str) -> str:
        """Clean and normalize raw extracted text"""
        if not text:
            return ""

        lines = text.split("\n")
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                cleaned_lines.append("")
                continue

            cleaned_line = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", line)
            cleaned_line = re.sub(r"\s+", " ", cleaned_line)

            if cleaned_line:
                cleaned_lines.append(cleaned_line)

        result = "\n".join(cleaned_lines)
        result = re.sub(r"\n{3,}", "\n\n", result)

        return result

    def _extract_name(self) -> Optional[str]:
        """Extract name - usually first non-empty line that looks like a name"""
        lines = self.cleaned_text.split("\n")

        for i, line in enumerate(lines[:5]):
            line = line.strip()
            if not line:
                continue

            line_clean = re.sub(r"[^\w\s\-\.\']", "", line)

            words = line_clean.split()
            if 2 <= len(words) <= 4:
                all_words_valid = all(
                    w[0].isupper() or w in ["de", "van", "von", "la", "le", "O'"]
                    for w in words
                    if len(w) > 1
                )
                line_lower = line.lower()
                if all_words_valid and not any(
                    c in line_lower
                    for c in ["@", "http", "www", "phone", "email", "address"]
                ):
                    return line

        return None

    def _extract_contact_info(
        self,
    ) -> Tuple[Optional[str], Optional[str], Optional[str], List[Profile]]:
        """Extract email, phone, location, and profiles"""
        email = None
        phone = None
        location = None
        profiles: List[Profile] = []

        text_lower = self.cleaned_text.lower()

        email_match = re.search(
            r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})", self.cleaned_text
        )
        if email_match:
            email = email_match.group(1)

        phone_match = re.search(
            r"[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}",
            self.cleaned_text,
        )
        if phone_match:
            phone = phone_match.group()

        location_patterns = [
            r"([A-Z][a-z]+,\s*[A-Z]{2}\s*\d{5})",
            r"([A-Z][a-zA-Z\s]+,\s*[A-Z][a-zA-Z]+)",
            r"(remote)",
        ]
        for pattern in location_patterns:
            loc_match = re.search(pattern, self.cleaned_text)
            if loc_match:
                location = loc_match.group(1)
                break

        linkedin_match = re.search(
            r"linkedin\.com/in/([a-zA-Z0-9_-]+)", self.cleaned_text
        )
        if linkedin_match:
            profiles.append(
                Profile(
                    network="LinkedIn",
                    username=linkedin_match.group(1),
                    url=f"https://linkedin.com/in/{linkedin_match.group(1)}",
                )
            )

        github_match = re.search(r"github\.com/([a-zA-Z0-9_-]+)", self.cleaned_text)
        if github_match:
            profiles.append(
                Profile(
                    network="GitHub",
                    username=github_match.group(1),
                    url=f"https://github.com/{github_match.group(1)}",
                )
            )

        return email, phone, location, profiles

    def _extract_summary(self) -> Optional[str]:
        """Extract professional summary"""
        patterns = [
            r"(?:summary|profile|objective|about me)[:\s]*\n(.*?)(?=\n\n[A-Z]|\n{2,}|\Z)",
            r"^(?!.*(?:experience|education|skills|projects)).*$",
        ]

        for pattern in patterns:
            match = re.search(pattern, self.cleaned_text, re.DOTALL | re.IGNORECASE)
            if match:
                text = match.group(1) if "(" in pattern else match.group(0)
                text = text.strip()
                if len(text) > 50 and len(text) < 1000:
                    return text

        return None

    def _extract_section(self, section_name: str) -> List[str]:
        """Extract a section's content"""
        patterns = self.SECTION_PATTERNS.get(section_name, [])

        for pattern in patterns:
            match = re.search(pattern, self.cleaned_text, re.IGNORECASE | re.DOTALL)
            if match:
                return [m.strip() for m in match.groups() if m.strip()]

        return []

    def _parse_experience(self) -> List[WorkExperience]:
        """Parse work experience section"""
        experiences: List[WorkExperience] = []

        exp_pattern = r"([A-Z][^0-9\n]+(?:Inc|LLC|Corp|Ltd|Company|Co\.|Technologies|Technologies?|Solutions|Software| Labs?|Studio|Group|Services)?)\s*[\|\n]\s*([^\n]+)\s*[\|\n]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})\s*[-–]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|Present|Current)"
        matches = re.finditer(exp_pattern, self.cleaned_text, re.IGNORECASE)

        for match in matches:
            company = match.group(1).strip()
            position = match.group(2).strip()
            start_date = match.group(3).strip()
            end_date = match.group(4).strip()

            start_idx = match.end()
            next_section = re.search(
                r"\n\n[A-Z][a-z]+[:\s]", self.cleaned_text[start_idx:]
            )
            end_idx = (
                start_idx + next_section.start()
                if next_section
                else len(self.cleaned_text)
            )

            section_text = self.cleaned_text[start_idx:end_idx]
            bullets = [b.strip() for b in re.findall(r"[-•*]\s*(.+)", section_text)]

            experiences.append(
                WorkExperience(
                    name=company,
                    position=position,
                    start_date=self._normalize_date(start_date),
                    end_date=self._normalize_date(end_date)
                    if end_date.lower() not in ["present", "current"]
                    else None,
                    highlights=bullets,
                )
            )

        if not experiences:
            simple_pattern = r"(?:^|\n)([^\n]+(?:Technologies?|Solutions|Software|Inc|LLC|Corp|Ltd|Company|Co\.))\s*[\n-]\s*([^\n]+?)(?=\n\n|\Z)"
            matches = re.finditer(simple_pattern, self.cleaned_text, re.MULTILINE)
            for match in matches:
                experiences.append(
                    WorkExperience(
                        name=match.group(1).strip(),
                        position=match.group(2).strip() if match.group(2) else "",
                    )
                )

        return experiences

    def _parse_education(self) -> List[Education]:
        """Parse education section"""
        education: List[Education] = []

        edu_pattern = r"((?:Bachelor|Master|PhD|B\.S\.|M\.S\.|B\.A\.|M\.A\.|MBA|Associate)[^\n]+)\s*[-–|\n]\s*([^\n]+?)(?:\s*[-–|\n]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}))?"
        matches = re.finditer(edu_pattern, self.cleaned_text, re.IGNORECASE)

        for match in matches:
            degree = match.group(1).strip()
            institution = match.group(2).strip()
            year = match.group(3).strip() if match.group(3) else None

            education.append(
                Education(
                    institution=institution,
                    area=degree,
                    study_type=degree,
                    start_date=year,
                    end_date=year,
                )
            )

        return education

    def _parse_skills(self) -> List[Skill]:
        """Parse skills section"""
        skills: List[Skill] = []

        skills_match = re.search(
            r"(?:skills|technologies|technical skills|competencies)[:\s]*\n(.*?)(?=\n\n[A-Z]|\n{2,}|\Z)",
            self.cleaned_text,
            re.IGNORECASE | re.DOTALL,
        )

        if skills_match:
            skills_text = skills_match.group(1)

            categories = re.split(r"\n(?=[A-Z][a-z]+[:\s])", skills_text)
            for category in categories:
                lines = category.strip().split("\n")
                if not lines:
                    continue

                if ":" in lines[0]:
                    parts = lines[0].split(":", 1)
                    category_name = parts[0].strip()
                    keywords = [k.strip() for k in parts[1].split(",") if k.strip()]
                    keywords.extend([k.strip() for k in lines[1:] if k.strip()])
                else:
                    category_name = "Technical Skills"
                    keywords = [
                        k.strip().lstrip("-•*").strip() for k in lines if k.strip()
                    ]

                keywords = [k for k in keywords if len(k) > 1 and len(k) < 50]

                if keywords:
                    skills.append(Skill(name=category_name, keywords=keywords))

        if not skills:
            all_skills = re.findall(r"([A-Za-z+#][\w+.#-]{2,30})", self.cleaned_text)
            tech_keywords = {
                "python",
                "javascript",
                "typescript",
                "java",
                "c++",
                "c#",
                "go",
                "rust",
                "react",
                "vue",
                "angular",
                "node",
                "django",
                "flask",
                "fastapi",
                "aws",
                "azure",
                "gcp",
                "docker",
                "kubernetes",
                "sql",
                "postgresql",
                "mongodb",
                "redis",
                "elasticsearch",
                "git",
                "linux",
                "tensorflow",
                "pytorch",
                "scikit",
                "pandas",
                "numpy",
                "html",
                "css",
                "graphql",
                "rest",
                "api",
                "microservices",
                "ci/cd",
                "agile",
                "scrum",
            }
            found_skills = [s for s in all_skills if s.lower() in tech_keywords]
            if found_skills:
                skills.append(
                    Skill(name="Technical Skills", keywords=list(set(found_skills)))
                )

        return skills

    def _parse_projects(self) -> List[Project]:
        """Parse projects section"""
        projects: List[Project] = []

        project_pattern = r"\*\*([^*]+)\*\*[:\s]*(.+?)(?=\n\n|\Z)"
        matches = re.finditer(project_pattern, self.cleaned_text, re.DOTALL)

        for match in matches:
            name = match.group(1).strip()
            desc = match.group(2).strip()

            bullets = re.findall(r"[-•*]\s*(.+)", desc)

            projects.append(
                Project(
                    name=name,
                    description=bullets[0] if bullets else desc,
                    highlights=bullets[1:] if len(bullets) > 1 else [],
                )
            )

        return projects

    def _parse_certifications(self) -> List[Certificate]:
        """Parse certifications section"""
        certs: List[Certificate] = []

        cert_pattern = r"([A-Z][^:\n]+(?:Certified|Certification|Certificate)[^:\n]*?)[:\s-]*([^\n]+)?"
        matches = re.finditer(cert_pattern, self.cleaned_text, re.IGNORECASE)

        for match in matches:
            name = match.group(1).strip()
            issuer = match.group(2).strip() if match.group(2) else "Unknown"

            year_match = re.search(r"\b(19|20)\d{2}\b", name)
            year = year_match.group() if year_match else None

            certs.append(Certificate(name=name, issuer=issuer, date=year))

        return certs

    def _normalize_date(self, date_str: str) -> str:
        """Normalize date to YYYY-MM format"""
        date_str = date_str.strip()

        month_map = {
            "jan": "01",
            "feb": "02",
            "mar": "03",
            "apr": "04",
            "may": "05",
            "jun": "06",
            "jul": "07",
            "aug": "08",
            "sep": "09",
            "oct": "10",
            "nov": "11",
            "dec": "12",
            "january": "01",
            "february": "02",
            "march": "03",
            "april": "04",
            "june": "06",
            "july": "07",
            "august": "08",
            "september": "09",
            "october": "10",
            "november": "11",
            "december": "12",
        }

        match = re.search(
            r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*(\d{4})",
            date_str,
            re.IGNORECASE,
        )
        if match:
            month = month_map.get(match.group(1)[:3].lower(), "01")
            year = match.group(2)
            return f"{year}-{month}"

        year_match = re.search(r"\b(19|20)\d{2}\b", date_str)
        if year_match:
            return f"{year_match.group()}-01"

        return date_str

    def parse(self) -> ResumeSchema:
        """
        Parse raw resume text into structured ResumeSchema.
        Returns a complete ResumeSchema object.
        """
        name = self._extract_name()
        if not name:
            name = "Candidate"

        email, phone, location, profiles = self._extract_contact_info()

        summary = self._extract_summary()

        work = self._parse_experience()
        education = self._parse_education()
        skills = self._parse_skills()
        projects = self._parse_projects()
        certifications = self._parse_certifications()

        basics = ResumeBasics(
            name=name,
            email=email,
            phone=phone,
            summary=summary,
            location=Location(address=location) if location else None,
            profiles=profiles,
        )

        return ResumeSchema(
            basics=basics,
            work=work,
            education=education,
            skills=skills,
            projects=projects,
            certificates=certifications,
        )

    def get_all_text(self) -> str:
        """Return cleaned raw text for LLM processing"""
        return self.cleaned_text


def extract_text_from_file(file_path: str, file_extension: str) -> str:
    """
    Extract text from PDF or DOCX file.

    Args:
        file_path: Path to the file
        file_extension: File extension ('.pdf' or '.docx')

    Returns:
        Extracted text as string
    """
    import os

    if file_extension.lower() == ".pdf":
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            return "\n".join(text_parts)
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")

    elif file_extension.lower() in [".docx", ".doc"]:
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n".join(text_parts)
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    else:
        raise ValueError(f"Unsupported file type: {file_extension}")
